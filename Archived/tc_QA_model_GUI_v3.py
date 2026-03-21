import threading
import time
import json
import csv
import os
import shutil
from typing import List, Optional, Callable, Dict

try:
    # Update import to whichever OpenAI client you're using
    from openai import OpenAI
except Exception:
    # Placeholder - user must install openai package or adapt to their client lib.
    OpenAI = None

# Optional libraries for file parsing / voice. All are optional and handled gracefully.
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    import speech_recognition as sr
except Exception:
    sr = None

try:
    import pyttsx3
except Exception:
    pyttsx3 = None

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    from tkinter.scrolledtext import ScrolledText
except Exception as e:
    raise RuntimeError("Tkinter is required to run this GUI script.") from e


def ask_questions_with_context(
    api_key: str,
    questions: List[str],
    context: Optional[str] = None,
    model: str = "gpt-4o",
    temperature: float = 0.3,
    save_path: Optional[str] = None,
    save_format: str = "json",      # "json", "txt", or "csv"
    colorize: bool = True,
    show_progress: bool = True,
    append_to_file: bool = False,
    max_retries: int = 2,
    retry_delay: float = 2.0,
    progress_callback: Optional[Callable[[str, str], None]] = None,
    provider: str = "openai",        # "openai", "mistral", "llama"
    stop_event: Optional[threading.Event] = None  # NEW: cancellation signal
) -> List[str]:
    """
    Ask multiple questions to Chat APIs (OpenAI supported inline). This function:
      - Appends prior Q&A as context for subsequent questions.
      - Optionally saves results to a file in json/txt/csv formats.
      - Reports progress via progress_callback(type, text) where type is 'q', 'a', or 'info'.
      - provider: select provider; currently only OpenAI client usage is implemented inline.
        For other providers, the API key is exported to an environment variable so external SDKs
        or runtimes can pick it up (user must integrate provider SDK calls if desired).
      - stop_event: optional threading.Event; if set, the function will stop between questions / retries.
    """
    def cb(t: str, msg: str):
        if progress_callback:
            try:
                progress_callback(t, "" if msg is None else str(msg))
            except Exception:
                # Swallow UI callback errors so they don't break the worker
                pass

    # Normalize provider
    provider_norm = (provider or "openai").strip().lower()

    # Export API key into provider-specific env var so third-party libs can find it.
    if api_key:
        if provider_norm == "openai":
            os.environ["OPENAI_API_KEY"] = api_key
        elif provider_norm == "mistral":
            os.environ["MISTRAL_API_KEY"] = api_key
        elif provider_norm == "llama":
            os.environ["LLAMA_API_KEY"] = api_key
        else:
            # default fallback
            os.environ["API_KEY"] = api_key

    if provider_norm != "openai":
        cb("info", f"Provider '{provider_norm}' selected. API key exported to environment; to use this provider natively, integrate its SDK/client and call it where OpenAI client is used.")
    if provider_norm == "openai":
        if OpenAI is None:
            cb("info", "Error: OpenAI client not available. Install and configure the openai package.")
            raise RuntimeError("OpenAI client not available")

        client = OpenAI(api_key=api_key)
    else:
        if OpenAI is not None:
            cb("info", f"Provider '{provider_norm}' selected but falling back to OpenAI client (if you intended to use the other provider, integrate its SDK).")
            client = OpenAI(api_key=api_key)
        else:
            cb("info", f"Provider '{provider_norm}' selected but no supported client found locally. Please install and configure the provider SDK or use OpenAI.")
            raise RuntimeError(f"No client available for provider '{provider_norm}'")

    answers: List[str] = []

    running_context = (context.strip() if context else "").strip()
    prior_qa = ""

    def is_code_like(text: str) -> bool:
        if not text:
            return False
        return (
            "def " in text
            or "class " in text
            or any(sym in text for sym in [";", "{", "}", "#include", "import ", "console.log", "function "])
        )

    def _save_results(save_path_local: str, save_format_local: str, records: List[dict], append_flag: bool):
        fmt = save_format_local.lower()
        if fmt == "json":
            if append_flag and os.path.exists(save_path_local):
                with open(save_path_local, "r", encoding="utf-8") as f:
                    try:
                        existing = json.load(f)
                        if not isinstance(existing, list):
                            existing = [existing]
                    except Exception:
                        existing = []
                merged = existing + records
                with open(save_path_local, "w", encoding="utf-8") as f:
                    json.dump(merged, f, ensure_ascii=False, indent=2)
            else:
                with open(save_path_local, "w", encoding="utf-8") as f:
                    json.dump(records, f, ensure_ascii=False, indent=2)
        elif fmt == "txt":
            mode = "a" if append_flag else "w"
            with open(save_path_local, mode, encoding="utf-8") as f:
                for r in records:
                    f.write(f"Q: {r.get('question')}\nA:\n{r.get('answer')}\n\n---\n\n")
        elif fmt == "csv":
            mode = "a" if append_flag else "w"
            file_exists = os.path.exists(save_path_local)
            with open(save_path_local, mode, newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=["index", "question", "answer"])
                if not file_exists or not append_flag:
                    writer.writeheader()
                for r in records:
                    writer.writerow({
                        "index": r.get("index"),
                        "question": r.get("question"),
                        "answer": r.get("answer")
                    })
        else:
            raise ValueError(f"Unsupported save_format: {save_format_local}")

    canceled = False

    for idx, question in enumerate(questions, start=1):
        # Honor stop request before starting the next question
        if stop_event is not None and stop_event.is_set():
            cb("info", f"Canceled by user before sending Q{idx}.")
            canceled = True
            break

        combined_context_parts = []
        if running_context:
            combined_context_parts.append(running_context)
        if prior_qa:
            combined_context_parts.append("Previous Q&A:\n" + prior_qa.strip())
        combined_context = "\n\n".join(combined_context_parts).strip()

        code_like = is_code_like(combined_context)

        if combined_context:
            if code_like:
                user_content = (
                    f"The following contains code and prior discussion/context:\n\n"
                    f"{combined_context}\n\n"
                    f"Question: {question}\n\n"
                    "Please respond with two sections:\n"
                    "1. Explanation: Explain what the code does in simple terms.\n"
                    "2. Modified Code: Provide updated code (only the code block) or improvements based on the question.\n"
                )
            else:
                user_content = (
                    "Read the following context and previous Q&A, then answer the question.\n\n"
                    f"Context:\n{combined_context}\n\n"
                    f"Question:\n{question}\n\n"
                    "Answer:"
                )
        else:
            user_content = (
                "Answer the following question concisely.\n\n"
                f"Question:\n{question}\n\n"
                "Answer:"
            )

        messages = [
            {"role": "system", "content": "You are a helpful assistant and skilled code interpreter."},
            {"role": "user", "content": user_content}
        ]

        # Notify UI of question about to be sent
        if show_progress:
            cb("q", f"Q{idx}: {question}")

        response_text = None
        for attempt in range(max_retries + 1):
            # Honor stop request inside retries
            if stop_event is not None and stop_event.is_set():
                cb("info", f"Canceled by user during retries for Q{idx}.")
                canceled = True
                break
            try:
                # Note: This code uses the OpenAI client's chat completions signature implemented earlier.
                response = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    temperature=temperature
                )
                # Defensive access - ensure nested fields exist
                try:
                    response_text = response.choices[0].message.content.strip()
                except Exception:
                    # Fallback: convert whatever to string
                    response_text = str(response) if response is not None else ""
                break
            except Exception as e:
                if attempt < max_retries:
                    cb("info", f"Transient error calling API (attempt {attempt+1}/{max_retries}). Retrying in {retry_delay}s...")
                    # Wait but allow interruption
                    waited = 0.0
                    while waited < retry_delay:
                        if stop_event is not None and stop_event.is_set():
                            cb("info", f"Canceled by user during retry wait for Q{idx}.")
                            canceled = True
                            break
                        time.sleep(0.2)
                        waited += 0.2
                    if canceled:
                        break
                else:
                    cb("info", f"Error calling API: {e}")
                    raise
        if canceled:
            break

        answer = response_text or ""
        answers.append(answer)

        prior_qa += f"Q: {question}\nA: {answer}\n\n"

        if show_progress:
            cb("a", f"A{idx}: {answer}")

    # Auto-save if requested
    if save_path and not (stop_event is not None and stop_event.is_set()):
        records = []
        for i, (q, a) in enumerate(zip(questions, answers), start=1):
            records.append({"index": i, "question": q, "answer": a})
        try:
            _save_results(save_path, save_format, records, append_to_file)
            if show_progress:
                cb("info", f"Results saved to {save_path} (format={save_format})")
        except Exception as e:
            cb("info", f"Warning: Failed to save results to {save_path}: {e}")
    elif stop_event is not None and stop_event.is_set():
        cb("info", "Run was canceled; results were not saved (partial results may exist in memory).")

    return answers


# --- Utility functions for extracting text from attachments ---

# def extract_text_from_file(path: str, max_chars: int = 20000) -> str:
#     """
#     Try to extract text from common document/image formats using available libraries.
#     If libraries aren't available or extraction fails, return a small descriptive placeholder.
#     """
#     if not os.path.exists(path):
#         return f"[Missing file: {os.path.basename(path)}]"

#     try:
#         ext = os.path.splitext(path.lower())[1]
#     except Exception:
#         ext = os.path.splitext(path)[1].lower()

#     try:
#         if ext in (".pdf",) and PyPDF2 is not None:
#             try:
#                 reader = PyPDF2.PdfReader(path)
#                 texts = []
#                 for p in reader.pages:
#                     try:
#                         texts.append(p.extract_text() or "")
#                     except Exception:
#                         continue
#                 text = "\n".join(texts).strip()
#                 if not text:
#                     return f"[PDF attached: {os.path.basename(path)} (no extractable text)]"
#                 return text[:max_chars]
#             except Exception:
#                 return f"[PDF attached: {os.path.basename(path)} (extraction error)]"
#         elif ext in (".docx",) and docx is not None:
#             try:
#                 doc = docx.Document(path)
#                 text = "\n".join([p.text for p in doc.paragraphs if p.text])
#                 if not text:
#                     return f"[DOCX attached: {os.path.basename(path)} (no extractable text)]"
#                 return text[:max_chars]
#             except Exception:
#                 return f"[DOCX attached: {os.path.basename(path)} (extraction error)]"
#         elif ext in (".csv",) and pd is not None:
#             try:
#                 df = pd.read_csv(path, dtype=str, encoding="utf-8", engine="python")
#                 text = df.head(200).to_csv(index=False)
#                 return f"CSV {os.path.basename(path)} content (preview):\n{text[:max_chars]}"
#             except Exception:
#                 return f"[CSV attached: {os.path.basename(path)} (extraction error)]"
#         elif ext in (".xls", ".xlsx") and pd is not None:
#             try:
#                 df = pd.read_excel(path, sheet_name=0, dtype=str)
#                 text = df.head(200).to_csv(index=False)
#                 return f"Excel {os.path.basename(path)} content (preview):\n{text[:max_chars]}"
#             except Exception:
#                 return f"[Excel attached: {os.path.basename(path)} (extraction error)]"
#         elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif") and Image is not None and pytesseract is not None:
#             try:
#                 img = Image.open(path)
#                 text = pytesseract.image_to_string(img)
#                 if not text.strip():
#                     return f"[Image attached: {os.path.basename(path)} (no text detected)]"
#                 return f"Image {os.path.basename(path)} OCR text:\n{text[:max_chars]}"
#             except Exception:
#                 return f"[Image attached: {os.path.basename(path)} (OCR error)]"
#         else:
#             # Unknown extension or missing libraries: return placeholder mentioning file name
#             return f"[Attachment: {os.path.basename(path)}] (content not extracted - missing parser)"
#     except Exception:
#         return f"[Attachment: {os.path.basename(path)}] (error reading file)"


def extract_text_from_file(path: str, max_chars: int = 20000) -> str:
    """
    Improved PDF / document / image text extraction with multiple fallbacks:
      - For PDFs: try PyPDF2/pypdf, then pdfminer.six, then OCR via pdf2image+pytesseract (if available).
      - For DOCX, CSV, Excel, and images: behave as before, with clearer diagnostics.
    Returns a truncated string (max_chars) or a descriptive placeholder telling which parser was attempted/missing.
    """
    if not os.path.exists(path):
        return f"[Missing file: {os.path.basename(path)}]"

    ext = os.path.splitext(path)[1].lower()

    # Helper to safely truncate
    def _t(s: Optional[str]):
        if not s:
            return ""
        return s[:max_chars]

    # Try PDF path
    if ext == ".pdf":
        # 1) Try PyPDF2 / pypdf
        try:
            if PyPDF2 is not None:
                try:
                    reader = PyPDF2.PdfReader(path)
                    texts = []
                    for p in reader.pages:
                        try:
                            t = p.extract_text() or ""
                            texts.append(t)
                        except Exception:
                            continue
                    joined = "\n".join(texts).strip()
                    if joined:
                        return _t(joined)
                    # if no text found continue to other methods
                except Exception:
                    # Fall through to other parsers
                    pass
        except Exception:
            pass

        # 2) Try pdfminer.six if available (better at some PDFs)
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract_text  # type: ignore
            try:
                txt = pdfminer_extract_text(path) or ""
                if txt.strip():
                    return _t(txt)
            except Exception:
                pass
        except Exception:
            # pdfminer not installed
            pass

        # 3) If still no text, try OCR path: pdf2image -> pytesseract (for scanned PDFs)
        try:
            # import locally so the function still runs if these libs aren't installed
            from pdf2image import convert_from_path  # type: ignore
            have_pdf2image = True
        except Exception:
            have_pdf2image = False

        if have_pdf2image and pytesseract is not None and Image is not None:
            try:
                # convert first N pages to images to limit processing
                pages = convert_from_path(path, first_page=1, last_page=5)  # limit pages for speed; adjust if needed
                ocr_texts = []
                for img in pages:
                    try:
                        t = pytesseract.image_to_string(img) or ""
                        ocr_texts.append(t)
                    except Exception:
                        continue
                joined_ocr = "\n".join(ocr_texts).strip()
                if joined_ocr:
                    return _t(f"PDF OCR text (first pages):\n{joined_ocr}")
                else:
                    return f"[PDF attached: {os.path.basename(path)} (OCR returned no text; install/update Tesseract/poppler if needed)]"
            except Exception as e:
                # Provide a diagnostic message about OCR failure
                return f"[PDF attached: {os.path.basename(path)} (OCR error: {e})]"
        else:
            # Inform which libraries are missing so the user can install them
            missing = []
            if PyPDF2 is None:
                missing.append("PyPDF2/pypdf")
            # check pdfminer
            try:
                import pdfminer  # type: ignore
            except Exception:
                missing.append("pdfminer.six")
            if pytesseract is None or Image is None:
                missing.append("pytesseract + Pillow (for OCR)")
            try:
                import pdf2image  # type: ignore
            except Exception:
                missing.append("pdf2image + poppler (for PDF->image OCR)")
            missing_list = ", ".join(sorted(set(missing))) if missing else "no-parsers"
            return f"[PDF attached: {os.path.basename(path)}] (content not extracted - missing/parsers: {missing_list})"

    # Non-PDF handling follows original logic, with clearer messages
    if ext == ".docx":
        if docx is not None:
            try:
                doc = docx.Document(path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
                if not text:
                    return f"[DOCX attached: {os.path.basename(path)} (no extractable text)]"
                return _t(text)
            except Exception:
                return f"[DOCX attached: {os.path.basename(path)} (extraction error)]"
        else:
            return f"[DOCX attached: {os.path.basename(path)}] (parser missing: python-docx)"

    if ext == ".csv":
        if pd is not None:
            try:
                df = pd.read_csv(path, dtype=str, encoding="utf-8", engine="python")
                text = df.head(200).to_csv(index=False)
                return f"CSV {os.path.basename(path)} content (preview):\n{_t(text)}"
            except Exception:
                return f"[CSV attached: {os.path.basename(path)} (extraction error)]"
        else:
            return f"[CSV attached: {os.path.basename(path)}] (parser missing: pandas)"

    if ext in (".xls", ".xlsx"):
        if pd is not None:
            try:
                df = pd.read_excel(path, sheet_name=0, dtype=str)
                text = df.head(200).to_csv(index=False)
                return f"Excel {os.path.basename(path)} content (preview):\n{_t(text)}"
            except Exception:
                return f"[Excel attached: {os.path.basename(path)} (extraction error)]"
        else:
            return f"[Excel attached: {os.path.basename(path)}] (parser missing: pandas + openpyxl/xlrd)"

    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"):
        if Image is not None and pytesseract is not None:
            try:
                img = Image.open(path)
                text = pytesseract.image_to_string(img)
                if not text.strip():
                    return f"[Image attached: {os.path.basename(path)} (no text detected)]"
                return f"Image {os.path.basename(path)} OCR text:\n{_t(text)}"
            except Exception:
                return f"[Image attached: {os.path.basename(path)} (OCR error)]"
        else:
            missing = []
            if Image is None:
                missing.append("Pillow")
            if pytesseract is None:
                missing.append("pytesseract")
            return f"[Image attached: {os.path.basename(path)}] (content not extracted - missing: {', '.join(missing)})"

    # Generic fallback for unknown extensions
    return f"[Attachment: {os.path.basename(path)}] (content not extracted - unsupported file type or missing parser)"


# --- Tkinter GUI ---

class QAGui:
    # popular model options - can be edited/extended
    MODELS = [
        "gpt-5", "gpt-5-mini",
        "gpt-4o", "gpt-4o-mini", "gpt-4o-realtime-preview", "gpt-4o-code",
        "gpt-4o-small", "gpt-3.5-turbo",
        # Popular Mistral / LLaMA-style models (names as commonly referenced)
        "mistral-1", "mistral-1-small",
        "llama-2-70b-chat", "llama-2-13b-chat", "llama-2-7b-chat"
    ]
    LOCKED_TEMP_MODELS = {"gpt-5", "gpt-5-mini"}

    def __init__(self, root: tk.Tk):
        self.root = root
        root.title("Tilendra's QA Chatbot")
        # Colorful palette
        self.bg = "#1f2b38"
        self.panel_bg = "#253241"
        self.accent = "#2dd4bf"  # teal
        self.info_bg = "#2b2f3a"

        root.configure(bg=self.bg)

        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass
        style.configure("TFrame", background=self.panel_bg)
        style.configure("TLabel", background=self.panel_bg, foreground="white")
        style.configure("TButton", background=self.accent, foreground="black")
        style.configure("TEntry", foreground="black")
        style.configure("TCombobox", foreground="black")
        style.configure("TCheckbutton", background=self.panel_bg, foreground="white")

        frm = ttk.Frame(root, padding="8", style="TFrame")
        frm.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)

        # Row 0: Provider, API key, model (dropdown), temperature (slider)
        top_row = ttk.Frame(frm, style="TFrame")
        top_row.grid(row=0, column=0, sticky="ew", pady=(0,8))
        ttk.Label(top_row, text="Provider:").grid(row=0, column=0, sticky="w")
        self.provider_var = tk.StringVar(value="openai")
        # Use a simple provider list directly for provider combobox
        self.provider_box = ttk.Combobox(top_row, textvariable=self.provider_var, values=["openai", "mistral", "llama"], width=12, state="readonly")
        self.provider_box.grid(row=0, column=1, sticky="w", padx=(4,8))
        self.provider_box.bind("<<ComboboxSelected>>", lambda e: self.on_provider_change())

        ttk.Label(top_row, text="API Key:").grid(row=0, column=2, sticky="w")
        self.api_key_var = tk.StringVar(value="")
        # show key with bullets
        self.api_key_entry = ttk.Entry(top_row, textvariable=self.api_key_var, width=36, show="•")
        self.api_key_entry.grid(row=0, column=3, sticky="ew", padx=(4,8))

        ttk.Label(top_row, text="Model:").grid(row=0, column=4, sticky="w")
        self.model_var = tk.StringVar(value="gpt-4o-mini")
        self.model_box = ttk.Combobox(top_row, textvariable=self.model_var, values=self.MODELS, width=22, state="readonly")
        self.model_box.grid(row=0, column=5, sticky="w", padx=(4,0))
        self.model_box.bind("<<ComboboxSelected>>", lambda e: self.on_model_change())

        ttk.Label(top_row, text="Temp:").grid(row=0, column=6, sticky="w", padx=(8,0))
        # temperature slider
        self.temp_var = tk.DoubleVar(value=0.3)
        self.temp_scale = tk.Scale(top_row, variable=self.temp_var, from_=0.0, to=1.0, resolution=0.05,
                                   orient="horizontal", length=140, showvalue=False, bg=self.panel_bg, troughcolor="#3a4754",
                                   highlightthickness=0)
        self.temp_scale.grid(row=0, column=7, sticky="w", padx=(4,0))
        self.temp_label = ttk.Label(top_row, text=f"{self.temp_var.get():.2f}")
        self.temp_label.grid(row=0, column=8, sticky="w", padx=(6,0))
        # Update label when slider moves
        self.temp_var.trace_add("write", lambda *args: self.temp_label.configure(text=f"{self.temp_var.get():.2f}"))

        # Row 1: Context (multi-line)
        ttk.Label(frm, text="Context (text or code) - Optional:").grid(row=1, column=0, sticky="w")
        self.context_text = ScrolledText(frm, height=8, background="#0f1720", foreground="#d1d5db", insertbackground="white")
        self.context_text.grid(row=2, column=0, sticky="nsew", pady=(0,8))
        frm.rowconfigure(2, weight=0)

        # Attachments frame under context
        attach_row = ttk.Frame(frm, style="TFrame")
        attach_row.grid(row=3, column=0, sticky="ew", pady=(0,8))
        ttk.Label(attach_row, text="Attachments:").grid(row=0, column=0, sticky="w")
        self.attach_btn = ttk.Button(attach_row, text="Attach File(s)...", command=self.attach_files)
        self.attach_btn.grid(row=0, column=1, sticky="w", padx=(6,4))

        self.remove_attach_btn = ttk.Button(attach_row, text="Remove Selected", command=self.remove_selected_attachment)
        self.remove_attach_btn.grid(row=0, column=2, sticky="w", padx=(6,4))

        # Listbox to show attached files
        self.attach_listbox = tk.Listbox(attach_row, height=3, bg="#0f1720", fg="#d1d5db", selectbackground="#3a4754", activestyle="none")
        self.attach_listbox.grid(row=1, column=0, columnspan=6, sticky="ew", pady=(6,0))
        attach_row.columnconfigure(5, weight=1)

        self.attachments: List[str] = []

        # Row 4: Questions (one per line)
        ttk.Label(frm, text="Questions (one per line):").grid(row=4, column=0, sticky="w")
        self.questions_text = ScrolledText(frm, height=6, background="#0f1720", foreground="#d1d5db", insertbackground="white")
        self.questions_text.grid(row=5, column=0, sticky="nsew", pady=(0,8))
        frm.rowconfigure(5, weight=0)

        # Voice controls
        voice_row = ttk.Frame(frm, style="TFrame")
        voice_row.grid(row=6, column=0, sticky="ew", pady=(0,8))
        self.voice_input_var = tk.BooleanVar(value=False)
        self.voice_output_var = tk.BooleanVar(value=False)
        self.voice_input_check = ttk.Checkbutton(voice_row, text="Enable Voice Input", variable=self.voice_input_var)
        self.voice_input_check.grid(row=0, column=0, sticky="w")
        self.voice_output_check = ttk.Checkbutton(voice_row, text="Enable Voice Output", variable=self.voice_output_var)
        self.voice_output_check.grid(row=0, column=1, sticky="w", padx=(8,0))
        self.record_btn = ttk.Button(voice_row, text="Record Question", command=self.on_record_voice)
        self.record_btn.grid(row=0, column=2, sticky="w", padx=(12,0))

        # Row 7: Save options - with slider switch (0 = No Save, 1 = Save)
        save_row = ttk.Frame(frm, style="TFrame")
        save_row.grid(row=7, column=0, sticky="ew", pady=(0,8))
        ttk.Label(save_row, text="Save option:").grid(row=0, column=0, sticky="w")
        self.save_enable_var = tk.IntVar(value=0)
        # Slider switch - 0 or 1
        self.save_slider = tk.Scale(save_row, variable=self.save_enable_var, from_=0, to=1, orient="horizontal", length=80,
                                    showvalue=False, resolution=1, bg=self.panel_bg, troughcolor="#3a4754",
                                    highlightthickness=0, command=lambda v: self.on_save_toggle())
        self.save_slider.grid(row=0, column=1, sticky="w", padx=(4,4))
        self.save_state_label = ttk.Label(save_row, text="No Save")
        self.save_state_label.grid(row=0, column=2, sticky="w", padx=(6,8))

        ttk.Label(save_row, text="Save to file:").grid(row=0, column=3, sticky="w")
        self.save_path_var = tk.StringVar(value="")
        self.save_entry = ttk.Entry(save_row, textvariable=self.save_path_var, width=36)
        self.save_entry.grid(row=0, column=4, sticky="w", padx=(4,4))
        self.browse_btn = ttk.Button(save_row, text="Browse...", command=self.browse_save)
        self.browse_btn.grid(row=0, column=5, sticky="w")

        ttk.Label(save_row, text="Format:").grid(row=0, column=6, sticky="w", padx=(8,0))
        self.save_format_var = tk.StringVar(value="json")
        self.save_format_box = ttk.Combobox(save_row, textvariable=self.save_format_var, values=["json", "txt", "csv"], width=6, state="readonly")
        self.save_format_box.grid(row=0, column=7, sticky="w", padx=(4,0))

        self.append_var = tk.BooleanVar(value=False)
        self.append_check = ttk.Checkbutton(save_row, text="Append", variable=self.append_var)
        self.append_check.grid(row=0, column=8, sticky="w", padx=(8,0))

        # Initially disable save controls (No Save)
        self.set_save_controls_enabled(False)

        # Row 8: Buttons
        btn_row = ttk.Frame(frm, style="TFrame")
        btn_row.grid(row=8, column=0, sticky="ew", pady=(0,8))
        self.run_btn = ttk.Button(btn_row, text="Run", command=self.on_run)
        self.run_btn.grid(row=0, column=0, sticky="w")
        ttk.Button(btn_row, text="Clear Output", command=self.clear_output).grid(row=0, column=1, sticky="w", padx=(8,0))
        ttk.Button(btn_row, text="Copy Output", command=self.copy_output).grid(row=0, column=2, sticky="w", padx=(8,0))
        ttk.Button(btn_row, text="Exit", command=root.quit).grid(row=0, column=3, sticky="w", padx=(8,0))

        # NEW: Stop button to cancel running worker
        self.stop_btn = ttk.Button(btn_row, text="Stop", command=self.on_stop)
        self.stop_btn.grid(row=0, column=4, sticky="w", padx=(8,0))
        try:
            self.stop_btn.configure(state="disabled")
        except Exception:
            pass

        # Row 9: Output text (color via tags)
        ttk.Label(frm, text="Output:").grid(row=9, column=0, sticky="w")
        self.output = ScrolledText(frm, height=18, state="normal", background="#071026", foreground="#dbeafe", insertbackground="white")
        self.output.grid(row=10, column=0, sticky="nsew")
        frm.rowconfigure(10, weight=1)

        # Text tags for colorizing
        self.output.tag_config("q", foreground="#7dd3fc", font=("TkDefaultFont", 10, "bold"))
        self.output.tag_config("a", foreground="#86efac", font=("TkDefaultFont", 10))
        self.output.tag_config("info", foreground="#ffd980", font=("TkDefaultFont", 9, "italic"))
        self.output.tag_config("error", foreground="#ff7b7b", font=("TkDefaultFont", 10, "bold"))

        # Worker thread handle
        self.worker_thread: Optional[threading.Thread] = None

        # Cancellation event for stopping the worker
        self.stop_event: Optional[threading.Event] = None

        # TTS engine (lazy init)
        self.tts_engine = None

        # Ensure temperature locked appropriately for initial model
        self.on_model_change()
        # Ensure provider label updated if needed
        self.on_provider_change()

    # --- Attachment handling ---
    def attach_files(self):
        try:
            # Use space-separated patterns for filetypes (widely supported by Tk)
            files = filedialog.askopenfilenames(title="Select files to attach",
                                                filetypes=[
                                                    ("All files", "*.*"),
                                                    ("PDF", "*.pdf"),
                                                    ("Word", "*.docx"),
                                                    ("CSV", "*.csv"),
                                                    ("Excel", "*.xls *.xlsx"),
                                                    ("Images", "*.png *.jpg *.jpeg *.bmp *.tiff *.gif")
                                                ])
        except Exception as e:
            # Defensive handling so dialog errors don't crash the app
            self.append_output("error", f"Failed to open file dialog: {e}")
            return

        if not files:
            # No files selected; be quiet
            return

        added = 0
        for f in files:
            # Some platforms may return bytes or None - ensure we only accept strings
            try:
                if not f:
                    continue
                f_str = str(f)
            except Exception:
                continue
            if f_str and f_str not in self.attachments:
                self.attachments.append(f_str)
                try:
                    display_name = os.path.basename(f_str) or f_str
                    self.attach_listbox.insert("end", display_name)
                except Exception:
                    # If insert fails for some reason, skip adding to listbox but still store attachment
                    pass
                added += 1
        if added:
            self.append_output("info", f"Attached {added} file(s).")
        else:
            self.append_output("info", "No new attachments were added.")

    def remove_selected_attachment(self):
        sel = self.attach_listbox.curselection()
        if not sel:
            return
        # remove all selected indices
        for idx in reversed(sel):
            try:
                self.attach_listbox.delete(idx)
            except Exception:
                pass
            try:
                del self.attachments[idx]
            except Exception:
                pass
        self.append_output("info", "Removed selected attachment(s).")

    # --- Voice (record + playback) ---
    def on_record_voice(self):
        if not self.voice_input_var.get():
            self.append_output("info", "Voice input is disabled. Enable 'Voice Input' to record.")
            return
        if sr is None:
            self.append_output("info", "SpeechRecognition library not available. Install 'speechrecognition' and a microphone backend (pyaudio or sounddevice).")
            return

        def record_thread():
            r = sr.Recognizer()
            try:
                with sr.Microphone() as source:
                    self.append_output("info", "Recording... Please speak now.")
                    audio = r.listen(source, timeout=10, phrase_time_limit=60)
                    self.append_output("info", "Processing audio...")
                    try:
                        text = r.recognize_google(audio)
                        if text:
                            # Append recognized text as a new question line
                            current_text = self.questions_text.get("1.0", "end").strip()
                            append_prefix = "\n" if current_text else ""
                            self.questions_text.insert("end", append_prefix + text + "\n")
                            self.append_output("info", f"Recognized: {text}")
                        else:
                            self.append_output("info", "No speech was recognized.")
                    except sr.UnknownValueError:
                        self.append_output("info", "Could not understand audio.")
                    except sr.RequestError as e:
                        self.append_output("info", f"Speech recognition service error: {e}")
            except Exception as e:
                self.append_output("error", f"Recording failed: {e}")

        threading.Thread(target=record_thread, daemon=True).start()

    def init_tts_engine(self):
        if self.tts_engine is not None:
            return True
        if pyttsx3 is None:
            return False
        try:
            self.tts_engine = pyttsx3.init()
            return True
        except Exception:
            self.tts_engine = None
            return False

    def play_tts(self, text: str):
        if not text:
            return
        # If stop requested, don't start TTS
        if self.stop_event is not None and self.stop_event.is_set():
            return
        # Run TTS on background thread to avoid blocking UI
        def tts_thread():
            t = "" if text is None else str(text)
            # Check once more before speaking
            if self.stop_event is not None and self.stop_event.is_set():
                return
            if self.init_tts_engine():
                try:
                    self.tts_engine.say(t)
                    self.tts_engine.runAndWait()
                except Exception:
                    pass
            else:
                # Fallback - on macOS use 'say' command
                try:
                    if os.name == "posix" and shutil.which("say") is not None:
                        import subprocess
                        subprocess.call(["say", t])
                except Exception:
                    # Can't TTS; silently ignore
                    pass

        threading.Thread(target=tts_thread, daemon=True).start()

    # --- Save/Run helpers ---
    def browse_save(self):
        default_ext = f".{self.save_format_var.get() or 'json'}"
        f = filedialog.asksaveasfilename(defaultextension=default_ext,
                                         filetypes=[("All files", "*.*"), ("JSON", "*.json"), ("Text", "*.txt"), ("CSV", "*.csv")])
        if f:
            self.save_path_var.set(f)

    def set_save_controls_enabled(self, enabled: bool):
        state = "normal" if enabled else "disabled"
        # Entry and browse and format and append
        try:
            self.save_entry.configure(state=state)
            self.browse_btn.configure(state=state)
            self.save_format_box.configure(state="readonly" if enabled else "disabled")
            # Checkbutton state handling defensively
            try:
                if enabled:
                    self.append_check.state(["!disabled"])
                else:
                    self.append_check.state(["disabled"])
            except Exception:
                pass
        except Exception:
            pass

    def on_save_toggle(self):
        val = int(self.save_enable_var.get())
        if val:
            self.save_state_label.configure(text="Save to File", foreground=self.accent)
            self.set_save_controls_enabled(True)
        else:
            self.save_state_label.configure(text="No Save", foreground="white")
            self.set_save_controls_enabled(False)

    def on_model_change(self):
        model = (self.model_var.get() or "").strip()
        if model in self.LOCKED_TEMP_MODELS:
            # Lock to 1.0 and disable slider
            self.temp_var.set(1.0)
            try:
                self.temp_scale.configure(state="disabled")
            except Exception:
                pass
        else:
            # Default to 0.3 if the current value was previously 1.0 due to a locked model
            if self.temp_var.get() == 1.0:
                self.temp_var.set(0.3)
            try:
                self.temp_scale.configure(state="normal")
            except Exception:
                pass

    def on_provider_change(self):
        provider = (self.provider_var.get() or "openai").strip().lower()
        # Update API Key label hint depending on provider via info message
        self.append_output("info", f"Provider set to '{provider}'. Place the provider API key in the API Key field; it will be exported to the environment variable for that provider when you run.")

    def append_output(self, typ: str, text: str):
        # Insert into text widget; runs in main/UI thread.
        # typ can be 'q', 'a', 'info', 'error'
        if typ not in ("q", "a", "info", "error"):
            typ = "info"
        safe_text = "" if text is None else str(text)
        try:
            self.output.configure(state="normal")
            # ensure we don't pass None to insert
            self.output.insert("end", safe_text + "\n\n", typ)
            self.output.see("end")
            self.output.configure(state="disabled")
        except Exception:
            # Fallback: try inserting without tag (defensive)
            try:
                self.output.configure(state="normal")
                self.output.insert("end", safe_text + "\n\n")
                self.output.see("end")
                self.output.configure(state="disabled")
            except Exception:
                # Give up silently - avoid crashing UI thread
                pass

    def clear_output(self):
        try:
            self.output.configure(state="normal")
            self.output.delete("1.0", "end")
            self.output.configure(state="disabled")
        except Exception:
            pass

    def copy_output(self):
        # Copy the entire output to the clipboard
        try:
            # Temporarily make widget normal to get text reliably
            prev_state = self.output.cget("state")
            self.output.configure(state="normal")
            content = self.output.get("1.0", "end")
            # Make sure content is a string (never None)
            content = "" if content is None else str(content).strip()
            self.output.configure(state=prev_state)
            if content:
                try:
                    # Some platforms may raise if clipboard args are None; ensure str passed
                    self.root.clipboard_clear()
                    self.root.clipboard_append(content)
                    self.append_output("info", "Output copied to clipboard.")
                except Exception as e:
                    self.append_output("error", f"Failed to copy output to clipboard: {e}")
            else:
                self.append_output("info", "Output is empty; nothing copied.")
        except Exception as e:
            self.append_output("error", f"Failed to copy output: {e}")

    def on_stop(self):
        # Called when user clicks Stop
        if self.stop_event is None:
            self.append_output("info", "No running job to stop.")
            return
        if self.stop_event.is_set():
            self.append_output("info", "Stop already requested.")
            return
        try:
            self.stop_event.set()
            self.append_output("info", "Stop requested. Attempting to cancel the job...")
            # If TTS is running, attempt to stop it
            try:
                if self.tts_engine is not None:
                    try:
                        # pyttsx3 supports stop()
                        self.tts_engine.stop()
                    except Exception:
                        pass
            except Exception:
                pass
            try:
                self.stop_btn.configure(state="disabled")
            except Exception:
                pass
        except Exception as e:
            self.append_output("error", f"Failed to request stop: {e}")

    def on_run(self):
        if self.worker_thread and self.worker_thread.is_alive():
            messagebox.showinfo("Running", "A job is already running. Please wait for it to finish.")
            return

        api_key = (self.api_key_var.get() or "").strip()
        if not api_key:
            messagebox.showwarning("Missing API Key", "Please enter your API key.")
            return

        questions_raw = (self.questions_text.get("1.0", "end") or "").strip()
        if not questions_raw:
            messagebox.showwarning("No Questions", "Please enter at least one question (one per line).")
            return

        questions = [q.strip() for q in questions_raw.splitlines() if q.strip()]
        context = (self.context_text.get("1.0", "end") or "").strip()
        model = self.model_var.get().strip() or "gpt-4o"
        provider = (self.provider_var.get() or "openai").strip().lower()
        try:
            temperature = float(self.temp_var.get())
        except Exception:
            temperature = 0.3

        # If save slider indicates save enabled, use path; otherwise no save
        save_enabled = bool(self.save_enable_var.get())
        save_path = (self.save_path_var.get() or "").strip() or None
        if not save_enabled:
            save_path = None

        save_format = (self.save_format_var.get() or "json").strip().lower() or "json"
        append_to_file = bool(self.append_var.get())

        # Disable run button while working
        try:
            self.run_btn.configure(state="disabled")
        except Exception:
            pass
        try:
            self.stop_btn.configure(state="normal")
        except Exception:
            pass
        self.append_output("info", "Starting job...")

        # Prepare combined context: context text + extracted text from attachments (best-effort)
        attachment_texts = []
        if self.attachments:
            self.append_output("info", f"Preparing {len(self.attachments)} attachment(s) for context (extraction may require optional libraries).")
            for p in self.attachments:
                try:
                    t = extract_text_from_file(p)
                    attachment_texts.append(f"Attachment - {os.path.basename(p)}:\n{t}")
                except Exception as e:
                    attachment_texts.append(f"Attachment - {os.path.basename(p)}: [error extracting: {e}]")

        combined_context_str = context or ""
        if attachment_texts:
            if combined_context_str:
                combined_context_str += "\n\n"
            combined_context_str += "\n\n".join(attachment_texts)

        def progress_cb(typ: str, msg: str):
            # This may be called from worker thread; marshal to main thread using after
            safe_msg = "" if msg is None else str(msg)
            self.root.after(0, self.append_output, typ if typ in ("q", "a", "info") else "info", safe_msg)
            # If voice output enabled and we get an answer, speak it
            if typ == "a" and self.voice_output_var.get():
                # Launch TTS in background
                # Trim long responses to a reasonable length for TTS
                speak_text = safe_msg
                # strip "A1: " prefix if present
                if ":" in speak_text and speak_text.strip().split(":", 1)[0].startswith("A"):
                    # remove leading label
                    speak_text = speak_text.split(":", 1)[1].strip()
                # Limit size
                if len(speak_text) > 4000:
                    speak_text = speak_text[:4000] + " ... (truncated)"
                # Start TTS
                self.play_tts(speak_text)

        def worker():
            try:
                # Create or clear stop event for this run
                self.stop_event = threading.Event()

                # Export provider API key to environment (ask_questions_with_context will do the export too)
                if api_key:
                    if provider == "openai":
                        os.environ["OPENAI_API_KEY"] = api_key
                    elif provider == "mistral":
                        os.environ["MISTRAL_API_KEY"] = api_key
                    elif provider == "llama":
                        os.environ["LLAMA_API_KEY"] = api_key
                    else:
                        os.environ["API_KEY"] = api_key

                # Call ask_questions_with_context (runs API calls)
                ask_questions_with_context(
                    api_key=api_key,
                    questions=questions,
                    context=combined_context_str,
                    model=model,
                    temperature=temperature,
                    save_path=save_path,
                    save_format=save_format,
                    colorize=False,    # colorized output handled by GUI, not colorama
                    show_progress=True,
                    append_to_file=append_to_file,
                    max_retries=2,
                    retry_delay=2.0,
                    progress_callback=progress_cb,
                    provider=provider,
                    stop_event=self.stop_event  # pass cancellation event
                )

                if self.stop_event is not None and self.stop_event.is_set():
                    self.root.after(0, self.append_output, "info", "Job canceled by user.")
                else:
                    self.root.after(0, self.append_output, "info", "Job completed.")
            except Exception as e:
                self.root.after(0, self.append_output, "error", f"Error during run: {e}")
            finally:
                # Re-enable run button in main thread and disable stop
                try:
                    self.root.after(0, lambda: self.run_btn.configure(state="normal"))
                    self.root.after(0, lambda: self.stop_btn.configure(state="disabled"))
                except Exception:
                    pass
                # Clear stop_event reference (run ended)
                try:
                    self.stop_event = None
                except Exception:
                    pass

        # Start worker thread
        self.worker_thread = threading.Thread(target=worker, daemon=True)
        self.worker_thread.start()


def main():
    root = tk.Tk()
    app = QAGui(root)
    root.geometry("1150x820")
    root.mainloop()


if __name__ == "__main__":
    main()



